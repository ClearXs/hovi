#!/usr/bin/env python3
"""
Markdown转Word文档转换器
将Markdown文件转换为格式良好的Word文档，支持：
- 标题层级转换
- 加粗、斜体格式化
- 表格转换
- 列表转换
- 自动生成目录
- 段落格式设置
"""

import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class MarkdownToWordConverter:
    def __init__(self, template_path=None):
        self.doc = Document(template_path) if template_path and os.path.exists(template_path) else Document()
        self.current_heading = None
        self.headings = []  # 用于生成目录
        self._setup_styles()

    def _setup_styles(self):
        """设置文档默认样式"""
        # 设置默认段落格式
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)

        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.5  # 1.5倍行距
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

    def convert_file(self, md_file_path, output_path=None):
        """转换单个Markdown文件到Word"""
        with open(md_file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        self._parse_markdown(content)

        if output_path:
            self.doc.save(output_path)
            print(f"已保存到: {output_path}")

        return self.doc

    def convert_directory(self, md_dir, output_path, add_toc=True):
        """转换目录下所有Markdown文件到一个Word文档"""
        md_files = sorted(Path(md_dir).rglob('*.md'))

        # 先收集所有标题用于生成目录
        for md_file in md_files:
            with open(md_file, 'r', encoding='utf-8') as f:
                content = f.read()
            self._collect_headings(content)

        # 再解析所有文件内容
        for md_file in md_files:
            with open(md_file, 'r', encoding='utf-8') as f:
                content = f.read()
            self._parse_markdown(content)

        # 添加目录
        if add_toc and self.headings:
            self._add_table_of_contents()

        self.doc.save(output_path)
        print(f"已保存到: {output_path}")
        return self.doc

    def _collect_headings(self, content):
        """收集所有标题用于生成目录"""
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('# '):
                self.headings.append((1, line[2:]))
            elif line.startswith('## '):
                self.headings.append((2, line[3:]))
            elif line.startswith('### '):
                self.headings.append((3, line[4:]))
            elif line.startswith('#### '):
                self.headings.append((4, line[5:]))
            elif line.startswith('##### '):
                self.headings.append((5, line[6:]))

    def _add_table_of_contents(self):
        """添加目录"""
        # 添加目录标题
        toc_heading = self.doc.add_heading('目  录', level=1)

        # 设置目录标题格式
        for run in toc_heading.runs:
            run.font.bold = True
            run.font.size = Pt(16)

        # 添加目录内容
        for level, text in self.headings:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

            # 根据层级设置缩进
            indent = (level - 1) * 0.5
            p.paragraph_format.left_indent = Cm(indent)

            # 添加标题文字
            run = p.add_run(text)
            run.font.size = Pt(12)

            # 添加页码占位符（Word会自动更新）
            # 使用制表符连接标题和页码
            p.add_run('\t')

    def _parse_markdown(self, content):
        """解析Markdown内容"""
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            # 跳过空行
            if not line:
                i += 1
                continue

            # 处理标题 (# 标题)
            if line.startswith('# '):
                self._add_heading(line[2:], 1)
            elif line.startswith('## '):
                self._add_heading(line[3:], 2)
            elif line.startswith('### '):
                self._add_heading(line[4:], 3)
            elif line.startswith('#### '):
                self._add_heading(line[5:], 4)
            elif line.startswith('##### '):
                self._add_heading(line[6:], 5)

            # 处理表格（更宽松的检测）
            elif line.strip().startswith('|') and '|' in line:
                table_lines = []
                # 收集连续的表格行
                while i < len(lines):
                    stripped = lines[i].strip()
                    if stripped.startswith('|') and '|' in stripped:
                        table_lines.append(stripped)
                        i += 1
                    elif stripped == '':
                        # 空行，表格结束
                        break
                    else:
                        break
                # 只有当有足够的表格行时才添加表格（至少2行：表头+数据）
                if len(table_lines) >= 2:
                    self._add_table(table_lines)
                    continue

            # 处理列表
            elif line.startswith('- ') or line.startswith('* '):
                self._add_list(line)
            elif re.match(r'^\d+\.\s', line):
                self._add_ordered_list(line)

            # 处理加粗和斜体
            else:
                self._add_paragraph(line)

            i += 1

    def _add_heading(self, text, level):
        """添加标题"""
        heading = self.doc.add_heading(text, level=level)
        # 设置标题格式
        for run in heading.runs:
            run.font.bold = True
            run.font.size = Pt(16 if level <= 2 else 14)

        # 设置段落格式
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

    def _add_paragraph(self, text):
        """添加段落，处理加粗和斜体"""
        if not text:
            return

        p = self.doc.add_paragraph()
        # 设置段落格式 - 首行缩进
        p.paragraph_format.first_line_indent = Cm(0.74)  # 2字符
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        self._format_text_with_markdown(p, text)

    def _format_text_with_markdown(self, paragraph, text):
        """格式化带有Markdown标记的文本（加粗、斜体等）"""

        # 如果没有特殊标记，直接添加文本
        if '**' not in text and '__' not in text:
            # 处理单独 * 和 _ 的情况
            if '*' not in text and '_' not in text:
                paragraph.add_run(text)
                return

        # 清理并处理加粗和斜体
        # 使用更全面的正则来处理所有情况

        # 分割文本为：普通文本、加粗文本、斜体文本
        # 匹配顺序：先匹配最长的

        # 使用 re.finditer 找到所有匹配
        pattern = r'(\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_)'
        matches = list(re.finditer(pattern, text))

        if not matches:
            # 没有匹配到任何格式标记，直接添加
            paragraph.add_run(text)
            return

        last_end = 0
        for match in matches:
            # 添加匹配之前的文本
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])

            matched = match.group(1)

            if matched.startswith('**') and matched.endswith('**'):
                # 加粗
                run = paragraph.add_run(matched[2:-2])
                run.font.bold = True
            elif matched.startswith('__') and matched.endswith('__'):
                # 加粗
                run = paragraph.add_run(matched[2:-2])
                run.font.bold = True
            elif matched.startswith('*') and matched.endswith('*'):
                # 斜体
                run = paragraph.add_run(matched[1:-1])
                run.font.italic = True
            elif matched.startswith('_') and matched.endswith('_'):
                # 斜体
                run = paragraph.add_run(matched[1:-1])
                run.font.italic = True
            else:
                # 普通文本
                paragraph.add_run(matched)

            last_end = match.end()

        # 添加剩余文本
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    def _add_list(self, text):
        """添加无序列表"""
        # 移除开头的 - 或 *
        clean_text = re.sub(r'^[\-\*]\s+', '', text)

        p = self.doc.add_paragraph(clean_text, style='List Bullet')
        # 设置段落格式
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = 1.5

    def _add_ordered_list(self, text):
        """添加有序列表"""
        # 移除开头的数字
        clean_text = re.sub(r'^\d+\.\s+', '', text)

        p = self.doc.add_paragraph(clean_text, style='List Number')
        # 设置段落格式
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = 1.5

    def _add_table(self, table_lines):
        """添加表格"""
        if len(table_lines) < 2:
            return

        # 解析表头
        headers = [h.strip() for h in table_lines[0].split('|')[1:-1]]

        if not headers:
            return

        # 创建表格
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # 添加表头
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)

        # 添加数据行 - 跳过表头，分隔符行需要跳过
        start_idx = 2 if len(table_lines) > 2 and '|---' in table_lines[1] else 1

        for line in table_lines[start_idx:]:
            if not line.strip() or '|---' in line:
                continue

            cells = [c.strip() for c in line.split('|')[1:-1]]
            if len(cells) != len(headers):
                continue

            row_cells = table.add_row().cells
            for i, cell in enumerate(cells):
                row_cells[i].text = cell
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10.5)


def main():
    import argparse

    parser = argparse.ArgumentParser(description='Markdown转Word文档')
    parser.add_argument('input', help='输入Markdown文件或目录')
    parser.add_argument('-o', '--output', help='输出Word文件路径')
    parser.add_argument('-t', '--template', help='Word模板文件')
    parser.add_argument('--no-toc', action='store_true', help='不生成目录')

    args = parser.parse_args()

    converter = MarkdownToWordConverter(args.template)
    add_toc = not args.no_toc

    if os.path.isdir(args.input):
        if not args.output:
            args.output = os.path.join(args.input, 'output.docx')
        converter.convert_directory(args.input, args.output, add_toc=add_toc)
    else:
        if not args.output:
            args.output = args.input.replace('.md', '.docx')
        converter.convert_file(args.input, args.output)


if __name__ == '__main__':
    main()
